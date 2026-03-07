"""Markdown to DOCX converter for CJK-heavy academic reports.

Uses markdown-it-py for parsing and python-docx for document generation.
Handles headings, tables, bold/italic, code blocks, blockquotes, and lists.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from markdown_it import MarkdownIt


def create_document(title: str) -> Document:
    """Create a new Document with CJK-friendly default styles."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    # Set default paragraph spacing
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    return doc


def add_heading(doc: Document, text: str, level: int):
    """Add a heading with appropriate level."""
    heading = doc.add_heading(text, level=min(level, 9))
    return heading


def parse_inline(text: str):
    """Parse inline markdown (bold, italic, code, links) and return segments."""
    segments = []
    # Pattern for: **bold**, *italic*, `code`, [text](url), plain text
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # bold
        r'|(\*(.+?)\*)'          # italic
        r'|(`(.+?)`)'            # inline code
        r'|(\[(.+?)\]\((.+?)\))' # link
        r'|([^*`\[]+)'           # plain text
    )
    for m in pattern.finditer(text):
        if m.group(2):  # bold
            segments.append(("bold", m.group(2)))
        elif m.group(4):  # italic
            segments.append(("italic", m.group(4)))
        elif m.group(6):  # code
            segments.append(("code", m.group(6)))
        elif m.group(8):  # link
            segments.append(("link", m.group(8), m.group(9)))
        elif m.group(10):  # plain
            segments.append(("plain", m.group(10)))
    return segments


def add_rich_paragraph(doc: Document, text: str, style=None, indent=False):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    segments = parse_inline(text)
    if not segments:
        p.add_run(text)
        return p
    for seg in segments:
        if seg[0] == "bold":
            run = p.add_run(seg[1])
            run.bold = True
        elif seg[0] == "italic":
            run = p.add_run(seg[1])
            run.italic = True
        elif seg[0] == "code":
            run = p.add_run(seg[1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
        elif seg[0] == "link":
            run = p.add_run(seg[1])
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            run.underline = True
        elif seg[0] == "plain":
            p.add_run(seg[1])
    return p


def add_table(doc: Document, rows: list[list[str]]):
    """Add a table from parsed rows (first row = header)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = cell_text.strip()
                if i == 0:  # header row bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    doc.add_paragraph()  # spacing after table


def convert_md_to_docx(md_path: str, docx_path: str, title: str = ""):
    """Convert a Markdown file to DOCX."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    doc = create_document(title)

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                # Flush table if pending
                if in_table:
                    add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows (---, :--:, etc.)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False
            # Don't increment i, reprocess current line

        # Horizontal rule
        if re.match(r'^---+\s*$', line.strip()):
            # Add a thin line / page break hint
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("─" * 60)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            quote_text = line.strip().lstrip("> ").strip()
            add_rich_paragraph(doc, quote_text, indent=True)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if ol_match:
            add_rich_paragraph(doc, line.strip(), style="List Number")
            i += 1
            continue

        # Unordered list
        if re.match(r'^[-*]\s+', line.strip()):
            text = re.sub(r'^[-*]\s+', '', line.strip())
            add_rich_paragraph(doc, text, style="List Bullet")
            i += 1
            continue

        # [TOC] - skip
        if line.strip() == "[TOC]":
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Caption/italics line (e.g., *表 1: ...*)
        caption_match = re.match(r'^\*(.+)\*$', line.strip())
        if caption_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption_match.group(1))
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, line.strip())
        i += 1

    # Flush remaining table
    if in_table:
        add_table(doc, table_rows)

    doc.save(docx_path)
    return Path(docx_path).stat().st_size


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx> [title]")
        sys.exit(1)
    md_file = sys.argv[1]
    docx_file = sys.argv[2]
    doc_title = sys.argv[3] if len(sys.argv) > 3 else ""
    size = convert_md_to_docx(md_file, docx_file, doc_title)
    print(f"OK: {docx_file} ({size:,} bytes)")
